"""
tests/test_parser.py
--------------------
Unit tests for extractor/parser.py.

All test documents are generated in-memory (no real resume files required).
PDF fixtures are created with PyMuPDF; DOCX fixtures with python-docx.
"""

from __future__ import annotations

import io
from pathlib import Path

import pymupdf as fitz  # PyMuPDF
import pytest
from docx import Document

from extractor.parser import (
    ParserError,
    extract_hyperlinks,
    extract_hyperlinks_from_docx,
    extract_hyperlinks_from_pdf,
    extract_text,
    extract_text_from_docx,
    extract_text_from_pdf,
)


# ---------------------------------------------------------------------------
# Fixture helpers — build minimal in-memory documents
# ---------------------------------------------------------------------------

def _make_pdf_bytes(pages: list[str]) -> bytes:
    """
    Create a minimal, text-selectable PDF in memory.

    Parameters
    ----------
    pages:
        A list of strings, one per page.

    Returns
    -------
    bytes
        Raw PDF bytes ready to be passed to ``extract_text_from_pdf``.
    """
    doc = fitz.open()
    for page_text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), page_text, fontsize=11)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """
    Create a minimal DOCX in memory.

    Parameters
    ----------
    paragraphs:
        A list of paragraph strings to add to the document body.

    Returns
    -------
    bytes
        Raw DOCX bytes ready to be passed to ``extract_text_from_docx``.
    """
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF extraction tests
# ---------------------------------------------------------------------------

class TestExtractTextFromPdf:
    """Tests for ``extract_text_from_pdf``."""

    def test_single_page_returns_text(self):
        """Text from a one-page PDF should be present in the output."""
        pdf_bytes = _make_pdf_bytes(["Alice Johnson\nalice@example.com\nPython Developer"])
        result = extract_text_from_pdf(pdf_bytes)
        assert "Alice Johnson" in result
        assert "alice@example.com" in result
        assert "Python Developer" in result

    def test_multi_page_returns_all_pages(self):
        """Text from every page should appear in the combined output."""
        pdf_bytes = _make_pdf_bytes(["Page one content", "Page two content"])
        result = extract_text_from_pdf(pdf_bytes)
        assert "Page one content" in result
        assert "Page two content" in result

    def test_returns_string(self):
        """Return type must be ``str``."""
        pdf_bytes = _make_pdf_bytes(["Some resume text"])
        result = extract_text_from_pdf(pdf_bytes)
        assert isinstance(result, str)

    def test_accepts_path(self, tmp_path: Path):
        """Parser should accept a ``Path`` object pointing to a PDF file."""
        pdf_file = tmp_path / "resume.pdf"
        pdf_file.write_bytes(_make_pdf_bytes(["Name: Bob Smith"]))
        result = extract_text_from_pdf(pdf_file)
        assert "Bob Smith" in result

    def test_accepts_file_like_object(self):
        """Parser should accept a binary ``BytesIO`` buffer."""
        buf = io.BytesIO(_make_pdf_bytes(["Skills: Python, SQL"]))
        result = extract_text_from_pdf(buf)
        assert "Python" in result

    def test_nonexistent_path_raises_parser_error(self, tmp_path: Path):
        """A path that does not exist should raise ``ParserError``."""
        with pytest.raises(ParserError, match="File not found"):
            extract_text_from_pdf(tmp_path / "missing.pdf")

    def test_invalid_bytes_raises_parser_error(self):
        """Corrupt/non-PDF bytes should raise ``ParserError``."""
        with pytest.raises(ParserError):
            extract_text_from_pdf(b"this is not a pdf")

    def test_page_separator_present_for_multi_page(self):
        """Multi-page PDFs should include a form-feed separator between pages."""
        pdf_bytes = _make_pdf_bytes(["First page", "Second page"])
        result = extract_text_from_pdf(pdf_bytes)
        assert "\f" in result


# ---------------------------------------------------------------------------
# DOCX extraction tests
# ---------------------------------------------------------------------------

class TestExtractTextFromDocx:
    """Tests for ``extract_text_from_docx``."""

    def test_basic_paragraphs_returned(self):
        """Text from all body paragraphs should appear in the output."""
        docx_bytes = _make_docx_bytes([
            "Jane Doe",
            "jane@example.com",
            "Experience: 3 years",
        ])
        result = extract_text_from_docx(docx_bytes)
        assert "Jane Doe" in result
        assert "jane@example.com" in result
        assert "Experience: 3 years" in result

    def test_returns_string(self):
        """Return type must be ``str``."""
        result = extract_text_from_docx(_make_docx_bytes(["Hello"]))
        assert isinstance(result, str)

    def test_accepts_path(self, tmp_path: Path):
        """Parser should accept a ``Path`` object pointing to a DOCX file."""
        docx_file = tmp_path / "resume.docx"
        docx_file.write_bytes(_make_docx_bytes(["Alice", "alice@mail.com"]))
        result = extract_text_from_docx(docx_file)
        assert "Alice" in result

    def test_accepts_file_like_object(self):
        """Parser should accept a binary ``BytesIO`` buffer."""
        buf = io.BytesIO(_make_docx_bytes(["Carlos", "Data Scientist"]))
        result = extract_text_from_docx(buf)
        assert "Carlos" in result

    def test_table_content_included(self):
        """Text inside DOCX tables should be included in the output."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Skills"
        table.cell(0, 1).text = "Python, Java, SQL"
        buf = io.BytesIO()
        doc.save(buf)
        result = extract_text_from_docx(buf.getvalue())
        assert "Python" in result
        assert "Skills" in result

    def test_invalid_bytes_raises_parser_error(self):
        """Non-DOCX bytes should raise ``ParserError``."""
        with pytest.raises(ParserError):
            extract_text_from_docx(b"not a docx file at all")

    def test_nonexistent_path_raises_parser_error(self, tmp_path: Path):
        """A path that does not exist should raise ``ParserError``."""
        with pytest.raises(ParserError, match="File not found"):
            extract_text_from_docx(tmp_path / "missing.docx")

    def test_paragraphs_newline_separated(self):
        """Each paragraph should appear on its own line."""
        docx_bytes = _make_docx_bytes(["Line one", "Line two", "Line three"])
        result = extract_text_from_docx(docx_bytes)
        lines = [l.strip() for l in result.splitlines() if l.strip()]
        assert "Line one" in lines
        assert "Line two" in lines
        assert "Line three" in lines


# ---------------------------------------------------------------------------
# Unified entry-point tests
# ---------------------------------------------------------------------------

class TestExtractText:
    """Tests for the unified ``extract_text`` dispatcher."""

    # --- PDF dispatch via filename hint ---

    def test_pdf_dispatched_by_filename(self):
        """A bytes source with filename='resume.pdf' should use the PDF parser."""
        pdf_bytes = _make_pdf_bytes(["Dispatch test PDF"])
        result = extract_text(pdf_bytes, filename="resume.pdf")
        assert "Dispatch test PDF" in result

    def test_pdf_dispatched_by_path(self, tmp_path: Path):
        """A ``Path`` source ending in .pdf should use the PDF parser."""
        pdf_file = tmp_path / "cv.pdf"
        pdf_file.write_bytes(_make_pdf_bytes(["Path dispatch PDF"]))
        result = extract_text(pdf_file)
        assert "Path dispatch PDF" in result

    # --- DOCX dispatch via filename hint ---

    def test_docx_dispatched_by_filename(self):
        """A bytes source with filename='resume.docx' should use the DOCX parser."""
        docx_bytes = _make_docx_bytes(["Dispatch test DOCX"])
        result = extract_text(docx_bytes, filename="resume.docx")
        assert "Dispatch test DOCX" in result

    def test_docx_dispatched_by_path(self, tmp_path: Path):
        """A ``Path`` source ending in .docx should use the DOCX parser."""
        docx_file = tmp_path / "cv.docx"
        docx_file.write_bytes(_make_docx_bytes(["Path dispatch DOCX"]))
        result = extract_text(docx_file)
        assert "Path dispatch DOCX" in result

    # --- Unsupported formats ---

    def test_unsupported_extension_raises_parser_error(self):
        """.txt uploads should raise ``ParserError`` with a helpful message."""
        with pytest.raises(ParserError, match="Unsupported file format"):
            extract_text(b"plain text content", filename="resume.txt")

    def test_unsupported_extension_doc_raises_parser_error(self):
        """Legacy .doc format is NOT supported and must raise ``ParserError``."""
        with pytest.raises(ParserError, match="Unsupported file format"):
            extract_text(b"", filename="resume.doc")

    def test_no_extension_raises_parser_error(self):
        """Bytes with no filename and no detectable extension raise ``ParserError``."""
        with pytest.raises(ParserError, match="Could not determine the file format"):
            extract_text(b"some bytes with no format hint")


    def test_case_insensitive_extension(self):
        """Extensions like .PDF or .Docx should be handled case-insensitively."""
        pdf_bytes = _make_pdf_bytes(["Case insensitive"])
        result = extract_text(pdf_bytes, filename="RESUME.PDF")
        assert "Case insensitive" in result

# ---------------------------------------------------------------------------
# Hyperlink extraction tests
# ---------------------------------------------------------------------------

_HYPERLINK_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _make_pdf_with_links(texts: list[str], links: list[str]) -> bytes:
    """Create a minimal PDF that embeds *links* as URI annotations."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "\n".join(texts), fontsize=12)
    y = 60
    for url in links:
        rect = fitz.Rect(50, y, 300, y + 20)
        page.insert_link({"kind": fitz.LINK_URI, "from": rect, "uri": url})
        y += 25
    result = doc.tobytes()
    doc.close()
    return result


def _make_docx_with_links(texts: list[str], links: list[str]) -> bytes:
    """Create a minimal DOCX that embeds *links* as hyperlink relationships."""
    from lxml import etree

    doc = Document()
    WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    for text, url in zip(texts, links):
        p = doc.add_paragraph()
        rel_id = doc.part.relate_to(url, _HYPERLINK_RELTYPE, is_external=True)
        hyperlink = etree.SubElement(p._p, f"{{{WORD_NS}}}hyperlink")
        hyperlink.set(f"{{{R_NS}}}id", rel_id)
        r = etree.SubElement(hyperlink, f"{{{WORD_NS}}}r")
        t = etree.SubElement(r, f"{{{WORD_NS}}}t")
        t.text = text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestHyperlinkExtraction:
    """Tests for extract_hyperlinks_from_pdf / _from_docx / extract_hyperlinks."""

    # -----------------------------------------------------------------------
    # PDF hyperlink extraction
    # -----------------------------------------------------------------------

    def test_pdf_linkedin_annotation(self):
        """LinkedIn URL embedded as PDF annotation is returned."""
        pdf = _make_pdf_with_links(
            ["LinkedIn"],
            ["https://www.linkedin.com/in/johndoe"],
        )
        result = extract_hyperlinks_from_pdf(pdf)
        assert "https://www.linkedin.com/in/johndoe" in result

    def test_pdf_github_annotation(self):
        """GitHub profile URL embedded as PDF annotation is returned."""
        pdf = _make_pdf_with_links(
            ["GitHub"],
            ["https://github.com/janedoe"],
        )
        result = extract_hyperlinks_from_pdf(pdf)
        assert "https://github.com/janedoe" in result

    def test_pdf_multiple_links(self):
        """Multiple annotation targets are all returned."""
        pdf = _make_pdf_with_links(
            ["LinkedIn", "GitHub"],
            ["https://linkedin.com/in/alice", "https://github.com/alice123"],
        )
        result = extract_hyperlinks_from_pdf(pdf)
        assert "https://linkedin.com/in/alice" in result
        assert "https://github.com/alice123" in result

    def test_pdf_no_links_returns_empty_list(self):
        """PDF with no annotation links returns an empty list."""
        pdf = _make_pdf_bytes(["No links here"])
        result = extract_hyperlinks_from_pdf(pdf)
        assert result == []

    def test_pdf_non_http_links_excluded(self):
        """mailto: and other non-HTTP URIs are excluded from results."""
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), "Email", fontsize=12)
        page.insert_link(
            {"kind": fitz.LINK_URI, "from": fitz.Rect(50, 60, 200, 80),
             "uri": "mailto:test@example.com"}
        )
        pdf = doc.tobytes()
        doc.close()
        result = extract_hyperlinks_from_pdf(pdf)
        assert result == []

    def test_pdf_deduplicates_same_url(self):
        """Duplicate annotation targets are de-duplicated."""
        url = "https://linkedin.com/in/johndoe"
        pdf = _make_pdf_with_links(["LinkedIn", "LinkedIn2"], [url, url])
        result = extract_hyperlinks_from_pdf(pdf)
        assert result.count(url) == 1

    def test_pdf_invalid_bytes_returns_empty_list(self):
        """Corrupt / non-PDF bytes return an empty list without raising."""
        result = extract_hyperlinks_from_pdf(b"not a pdf")
        assert result == []

    # -----------------------------------------------------------------------
    # DOCX hyperlink extraction
    # -----------------------------------------------------------------------

    def test_docx_linkedin_relationship(self):
        """LinkedIn URL stored as a DOCX relationship is returned."""
        docx = _make_docx_with_links(
            ["LinkedIn"],
            ["https://www.linkedin.com/in/johndoe"],
        )
        result = extract_hyperlinks_from_docx(docx)
        assert "https://www.linkedin.com/in/johndoe" in result

    def test_docx_github_relationship(self):
        """GitHub URL stored as a DOCX relationship is returned."""
        docx = _make_docx_with_links(
            ["GitHub"],
            ["https://github.com/janedoe"],
        )
        result = extract_hyperlinks_from_docx(docx)
        assert "https://github.com/janedoe" in result

    def test_docx_multiple_links(self):
        """Multiple DOCX hyperlink relationships are all returned."""
        docx = _make_docx_with_links(
            ["LinkedIn", "GitHub"],
            ["https://linkedin.com/in/bob", "https://github.com/bob99"],
        )
        result = extract_hyperlinks_from_docx(docx)
        assert "https://linkedin.com/in/bob" in result
        assert "https://github.com/bob99" in result

    def test_docx_no_hyperlinks_returns_empty_list(self):
        """A plain DOCX with no hyperlinks returns an empty list."""
        docx = _make_docx_bytes(["No links here"])
        result = extract_hyperlinks_from_docx(docx)
        assert result == []

    def test_docx_invalid_bytes_returns_empty_list(self):
        """Corrupt / non-DOCX bytes return an empty list without raising."""
        result = extract_hyperlinks_from_docx(b"not a docx")
        assert result == []

    # -----------------------------------------------------------------------
    # Unified extract_hyperlinks()
    # -----------------------------------------------------------------------

    def test_unified_pdf(self):
        """extract_hyperlinks() dispatches to PDF extractor."""
        pdf = _make_pdf_with_links(["Link"], ["https://linkedin.com/in/test"])
        result = extract_hyperlinks(pdf, filename="resume.pdf")
        assert "https://linkedin.com/in/test" in result

    def test_unified_docx(self):
        """extract_hyperlinks() dispatches to DOCX extractor."""
        docx = _make_docx_with_links(["Link"], ["https://github.com/testuser"])
        result = extract_hyperlinks(docx, filename="resume.docx")
        assert "https://github.com/testuser" in result

    def test_unified_unknown_format_returns_empty(self):
        """Unknown file extensions return an empty list without raising."""
        result = extract_hyperlinks(b"data", filename="resume.txt")
        assert result == []

    def test_unified_no_filename_returns_empty(self):
        """Without a filename hint, unknown bytes return an empty list."""
        result = extract_hyperlinks(b"data")
        assert result == []

