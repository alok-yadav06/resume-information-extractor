"""
parser.py
---------
Converts supported resume files (PDF, DOCX) into raw plain text.

This module is the *first* stage of the pipeline.  It does nothing more than
text extraction — no cleaning, no pattern matching, no field detection.  Keeping
parsing separate from extraction logic means each stage can be tested and improved
independently.

Supported formats
-----------------
- PDF  — extracted via PyMuPDF (``fitz``)
- DOCX — extracted via ``python-docx``

Public API
----------
extract_text_from_pdf(source)  -> str
extract_text_from_docx(source) -> str
extract_text(source, filename) -> str
"""

from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Union

import pymupdf as fitz  # PyMuPDF — use pymupdf alias to avoid deprecation warning
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# ---------------------------------------------------------------------------
# Type alias — a "source" is either a filesystem path or a binary buffer
# (e.g. the object returned by Streamlit's file_uploader).
# ---------------------------------------------------------------------------
FileSource = Union[str, Path, io.IOBase, bytes]


# ---------------------------------------------------------------------------
# Custom exception
# ---------------------------------------------------------------------------

class ParserError(Exception):
    """Raised when a resume file cannot be parsed."""


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _to_bytes(source: FileSource) -> bytes:
    """
    Normalise *source* to a raw ``bytes`` object so that both filesystem paths
    and in-memory buffers can be handled by a single code path.

    Parameters
    ----------
    source:
        A filesystem path (``str`` or ``Path``), a binary file-like object, or
        a raw ``bytes`` object.

    Returns
    -------
    bytes
        The full file content as bytes.

    Raises
    ------
    ParserError
        If the file does not exist or cannot be read.
    """
    if isinstance(source, bytes):
        return source

    if isinstance(source, (str, Path)):
        path = Path(source)
        if not path.exists():
            raise ParserError(f"File not found: {path}")
        try:
            return path.read_bytes()
        except OSError as exc:
            raise ParserError(f"Could not read file '{path}': {exc}") from exc

    # Assume file-like object (e.g. Streamlit UploadedFile, io.BytesIO, …)
    try:
        return source.read()
    except AttributeError as exc:
        raise ParserError(
            f"Unsupported source type '{type(source).__name__}'. "
            "Expected a file path, bytes, or a file-like object."
        ) from exc


def _detect_extension(source: FileSource, filename: str | None) -> str:
    """
    Return the lower-case file extension for *source*.

    Priority order:
    1. ``filename`` argument (useful for Streamlit uploads that keep the
       original name in ``UploadedFile.name``).
    2. The path string / Path object itself.
    3. Empty string if neither is available.
    """
    if filename:
        return Path(filename).suffix.lower()
    if isinstance(source, (str, Path)):
        return Path(source).suffix.lower()
    # For bare bytes / generic IO objects with no filename we cannot guess.
    return ""


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------

def extract_text_from_pdf(source: FileSource) -> str:
    """
    Extract all text from a PDF file using PyMuPDF.

    Text is extracted page by page and concatenated with a form-feed character
    (``\\f``) between pages so that callers can easily detect page boundaries
    if needed.  Within each page the natural reading order produced by
    PyMuPDF is preserved.

    Parameters
    ----------
    source:
        A filesystem path, raw bytes, or a binary file-like object pointing to
        a PDF file.

    Returns
    -------
    str
        The full plain-text content of the PDF.

    Raises
    ------
    ParserError
        If the PDF cannot be opened or if no text at all could be extracted
        (e.g. the file is a scanned image without an OCR layer).
    """
    raw = _to_bytes(source)

    try:
        # ``fitz.open`` accepts a ``bytes`` stream via the ``stream`` parameter.
        doc = fitz.open(stream=raw, filetype="pdf")
    except Exception as exc:
        raise ParserError(
            f"Could not open PDF file. "
            f"The file may be corrupted or password-protected. ({exc})"
        ) from exc

    pages_text: list[str] = []

    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            # "text" mode preserves word order left-to-right, top-to-bottom.
            page_text = page.get_text("text")
            if page_text.strip():
                pages_text.append(page_text)
    finally:
        doc.close()

    if not pages_text:
        raise ParserError(
            "No readable text was found in this PDF. "
            "The file may be a scanned image. "
            "Please upload a PDF with selectable text."
        )

    # Join pages; "\f" (form-feed) marks page breaks, "\n\n" gives a clear
    # visual gap when the text is later printed or inspected.
    return "\f\n".join(pages_text)


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------

def extract_text_from_docx(source: FileSource) -> str:
    """
    Extract all text from a DOCX file using python-docx.

    Text is collected from:
    - Body paragraphs (in document order).
    - Table cells (row by row, cell by cell), which are appended after all
      body paragraphs.

    Parameters
    ----------
    source:
        A filesystem path, raw bytes, or a binary file-like object pointing to
        a valid ``.docx`` file.

    Returns
    -------
    str
        The full plain-text content of the DOCX document.

    Raises
    ------
    ParserError
        If the file is not a valid DOCX document or cannot be read.
    """
    raw = _to_bytes(source)
    buffer = io.BytesIO(raw)

    try:
        doc = Document(buffer)
    except PackageNotFoundError as exc:
        raise ParserError(
            "The uploaded file does not appear to be a valid DOCX document. "
            "Please check the file and try again."
        ) from exc
    except Exception as exc:
        raise ParserError(
            f"Could not open DOCX file. The file may be corrupted. ({exc})"
        ) from exc

    lines: list[str] = []

    # --- Body paragraphs ---------------------------------------------------
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # --- Table cells -------------------------------------------------------
    # Resume tables often hold key-value pairs (e.g. "Skills | Python, Java").
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                lines.append("  |  ".join(row_cells))

    if not lines:
        raise ParserError(
            "No readable text was found in this DOCX file. "
            "The document may be empty."
        )

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Unified entry point
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_text(
    source: FileSource,
    filename: str | None = None,
) -> str:
    """
    Determine the file type and delegate to the appropriate parser.

    This is the primary function the rest of the pipeline should call.  It
    accepts both filesystem paths and Streamlit ``UploadedFile`` objects
    (or any binary file-like object) as long as a ``filename`` hint is
    provided for format detection when the source itself carries no extension.

    Parameters
    ----------
    source:
        A filesystem path, raw bytes, or a binary file-like object.
    filename:
        Optional original filename string used to determine the file type
        (e.g. ``uploaded_file.name`` from Streamlit).  If *source* is already
        a path, this parameter is ignored.

    Returns
    -------
    str
        The full plain-text content of the resume.

    Raises
    ------
    ParserError
        If the file format is unsupported or if parsing fails.
    """
    ext = _detect_extension(source, filename)

    if ext == ".pdf":
        return extract_text_from_pdf(source)

    if ext == ".docx":
        return extract_text_from_docx(source)

    if ext == "":
        raise ParserError(
            "Could not determine the file format. "
            "Please pass a filename hint (e.g. filename='resume.pdf') or "
            "use a file path directly."
        )

    raise ParserError(
        f"Unsupported file format '{ext}'. "
        "Please upload a PDF or DOCX resume."
    )
